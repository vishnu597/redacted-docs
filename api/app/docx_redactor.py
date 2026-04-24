from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from .detector import PrivacyFilterDetector
from .schemas import DetectedSpan, RedactionSummary
from .utils import build_summary, clamp_excerpt

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def redact_docx(
    input_path: Path,
    output_path: Path,
    detector: PrivacyFilterDetector,
) -> RedactionSummary:
    document = Document(str(input_path))
    summary_spans: list[DetectedSpan] = []
    cursor = 0

    for paragraph in iter_document_paragraphs(document):
        text = paragraph.text
        if not text.strip():
            cursor += len(text) + 2
            continue

        local_spans = detector.detect(text)
        if local_spans:
            redact_paragraph(paragraph, local_spans)
            summary_spans.extend(
                [
                    span.model_copy(
                        update={
                            "start": span.start + cursor,
                            "end": span.end + cursor,
                            "text": clamp_excerpt(span.text),
                        }
                    )
                    for span in local_spans
                ]
            )

        cursor += len(text) + 2

    document.save(str(output_path))
    return build_summary(summary_spans)


def iter_document_paragraphs(document: Document):
    yield from iter_story_paragraphs(document)

    seen_story_ids: set[int] = set()
    for section in document.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            story_id = id(story._element)
            if story_id in seen_story_ids:
                continue
            seen_story_ids.add(story_id)
            yield from iter_story_paragraphs(story)


def iter_story_paragraphs(story):
    for paragraph in story.paragraphs:
        yield paragraph
    for table in story.tables:
        yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell)


def iter_cell_paragraphs(cell: _Cell):
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        yield from iter_table_paragraphs(table)


def redact_paragraph(paragraph: Paragraph, spans: list[DetectedSpan]) -> None:
    if not paragraph.runs:
        return

    run_ranges = build_run_ranges(paragraph)
    segments: list[tuple[str, Run, bool]] = []

    for run, run_start, run_end in run_ranges:
        cursor = run_start
        overlapping_spans = [
            span for span in spans if span.end > run_start and span.start < run_end
        ]

        if not overlapping_spans:
            if run.text:
                segments.append((run.text, run, False))
            continue

        for span in overlapping_spans:
            if cursor < span.start:
                segments.append((paragraph.text[cursor:span.start], run, False))

            redaction_start = max(cursor, span.start)
            redaction_end = min(run_end, span.end)
            if redaction_end > redaction_start:
                segments.append(
                    (
                        redaction_blocks(paragraph.text[redaction_start:redaction_end]),
                        run,
                        True,
                    )
                )
            cursor = redaction_end

        if cursor < run_end:
            segments.append((paragraph.text[cursor:run_end], run, False))

    clear_paragraph_content(paragraph)

    for text, source_run, is_redacted in segments:
        if not text:
            continue
        target_run = paragraph.add_run(text)
        copy_run_style(target_run, source_run)
        if is_redacted:
            target_run.font.color.rgb = RGBColor(0, 0, 0)
            apply_black_shading(target_run)


def build_run_ranges(paragraph: Paragraph) -> list[tuple[Run, int, int]]:
    ranges: list[tuple[Run, int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        start = cursor
        cursor += len(run.text)
        ranges.append((run, start, cursor))
    return ranges


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def copy_run_style(target: Run, source: Run) -> None:
    if source.style is not None:
        target.style = source.style
    if source._r.rPr is not None:
        existing = target._r.find(qn("w:rPr"))
        if existing is not None:
            target._r.remove(existing)
        target._r.insert(0, deepcopy(source._r.rPr))


def redaction_blocks(text: str) -> str:
    return "".join("█" if char != "\t" else "\t" for char in text)


def apply_black_shading(run: Run) -> None:
    properties = run._r.get_or_add_rPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "000000")
